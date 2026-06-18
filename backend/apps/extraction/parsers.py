"""
File parsing utilities.

Downloads files from remote storage and extracts raw text
using specialized libraries.

OCR provider is controlled by the IMAGE_OCR_PROVIDER setting:
  - "gemini"       (default) — uses Gemini Vision, free tier, no billing required
  - "google_vision" — uses Google Cloud Vision API, requires billing enabled on GCP
"""

import io
import base64
import os
import tempfile
import requests
import pdfplumber
import docx
from django.conf import settings

from apps.files.models import FileType


def extract_text_from_url(url: str, file_type: str) -> str:
    """Download a file and extract its text content based on its type."""
    if not url:
        return ""

    if file_type == FileType.IMAGE:
        from django.conf import settings
        provider = getattr(settings, "IMAGE_OCR_PROVIDER", "gemini")
        if provider == "ocr_space":
            return _ocr_ocrspace_url(url)

    temp_path = _download_to_temp_file(url)
    try:
        if file_type == FileType.PDF:
            return _extract_from_pdf_path(temp_path)
        elif file_type == FileType.DOCX:
            return _extract_from_docx_path(temp_path)
        elif file_type == FileType.IMAGE:
            with open(temp_path, "rb") as file_obj:
                return _extract_from_image(file_obj.read())
    finally:
        try:
            os.unlink(temp_path)
        except OSError:
            pass

    return ""


def _download_to_temp_file(url: str) -> str:
    max_bytes = getattr(settings, "MAX_FILE_SIZE_BYTES", 10 * 1024 * 1024)
    timeout = getattr(settings, "FILE_DOWNLOAD_TIMEOUT_SECONDS", 30)

    with requests.get(url, stream=True, timeout=timeout) as response:
        response.raise_for_status()

        content_length = response.headers.get("content-length")
        if content_length:
            try:
                if int(content_length) > max_bytes:
                    raise RuntimeError(f"File exceeds maximum size of {max_bytes} bytes.")
            except ValueError:
                pass

        temp_file = tempfile.NamedTemporaryFile(delete=False)
        temp_path = temp_file.name
        bytes_written = 0

        try:
            with temp_file:
                for chunk in response.iter_content(chunk_size=1024 * 1024):
                    if not chunk:
                        continue
                    bytes_written += len(chunk)
                    if bytes_written > max_bytes:
                        raise RuntimeError(f"File exceeds maximum size of {max_bytes} bytes.")
                    temp_file.write(chunk)
        except Exception:
            try:
                os.unlink(temp_path)
            except OSError:
                pass
            raise

    return temp_path


def _extract_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n".join(pages)
    except Exception as e:
        raise RuntimeError(f"PDF extraction failed: {str(e)}")


def _extract_from_pdf_path(file_path: str) -> str:
    """Extract text from a PDF file path using pdfplumber."""
    try:
        with pdfplumber.open(file_path) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n".join(pages)
    except Exception as e:
        raise RuntimeError(f"PDF extraction failed: {str(e)}")


def _extract_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        raise RuntimeError(f"DOCX extraction failed: {str(e)}")


def _extract_from_docx_path(file_path: str) -> str:
    """Extract text from a DOCX file path using python-docx."""
    try:
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        raise RuntimeError(f"DOCX extraction failed: {str(e)}")


def _extract_from_image(file_bytes: bytes) -> str:
    """Route image OCR to the configured provider.

    Reads IMAGE_OCR_PROVIDER from Django settings:
      - "gemini"        → Gemini Vision (default, free tier)
      - "google_vision" → Google Cloud Vision API (requires billing)
    """
    from django.conf import settings
    provider = getattr(settings, "IMAGE_OCR_PROVIDER", "gemini")

    if provider == "google_vision":
        return _ocr_google_vision(file_bytes)
    return _ocr_gemini(file_bytes)


# ─── OCR Backends ─────────────────────────────────────────────────────────────

def _ocr_gemini(file_bytes: bytes) -> str:
    """Extract text from images using Gemini Vision (free tier).

    Uses Gemini's multimodal capability. No billing required beyond
    the standard Gemini API free tier.
    """
    try:
        from django.conf import settings
        import google.generativeai as genai  # type: ignore

        genai.configure(api_key=settings.GEMINI_API_KEY)  # type: ignore
        model = genai.GenerativeModel(settings.GEMINI_MODEL)  # type: ignore

        image_part = {
            "mime_type": "image/jpeg",
            "data": base64.b64encode(file_bytes).decode("utf-8"),
        }

        prompt = (
            "Extract ALL text from this image exactly as it appears. "
            "Return only the raw text content, no commentary or formatting. "
            "Preserve line breaks where they appear in the image."
        )

        response = model.generate_content([prompt, image_part])
        return response.text.strip() if response.text else ""
    except Exception as e:
        raise RuntimeError(f"Image OCR failed (Gemini): {str(e)}")


def _ocr_google_vision(file_bytes: bytes) -> str:
    """Extract text from images using Google Cloud Vision API.

    Requires:
      - Billing enabled on the GCP project
      - GOOGLE_VISION_CREDENTIALS set to the service account JSON path
    """
    try:
        from django.conf import settings
        from google.cloud import vision
        from google.oauth2 import service_account

        creds = None
        if getattr(settings, "GOOGLE_VISION_CREDENTIALS", ""):
            creds = service_account.Credentials.from_service_account_file(
                settings.GOOGLE_VISION_CREDENTIALS
            )

        client = vision.ImageAnnotatorClient(credentials=creds)
        image = vision.Image(content=file_bytes)

        response = client.document_text_detection(image=image)  # type: ignore
        if response.error.message:
            raise RuntimeError(f"Vision API error: {response.error.message}")

        if response.full_text_annotation:
            return response.full_text_annotation.text
        return ""
    except Exception as e:
        raise RuntimeError(f"Image OCR failed (Google Vision): {str(e)}")


def _ocr_ocrspace_url(file_url: str) -> str:
    """Extract text from images using OCR.space URL parsing."""
    import httpx
    from django.conf import settings

    try:
        with httpx.Client(timeout=30) as client:
            response = client.post(
                "https://api.ocr.space/parse/image",
                data={
                    "apikey": getattr(settings, "OCR_SPACE_API_KEY", ""),
                    "url": file_url,
                    "language": "eng",
                    "isOverlayRequired": False,
                    "OCREngine": 2,
                    "filetype": "Auto",
                },
            )
            response.raise_for_status()
    except httpx.TimeoutException:
        raise RuntimeError("OCR.space request timed out")
    except httpx.HTTPStatusError as e:
        raise RuntimeError(f"OCR.space HTTP error: {e.response.status_code}")

    result = response.json()

    if result.get("IsErroredOnProcessing"):
        raise RuntimeError(
            result.get("ErrorMessage")
            or result.get("ErrorDetails")
            or "OCR.space processing failed"
        )

    parsed = result.get("ParsedResults", [])
    if not parsed:
        raise RuntimeError("OCR.space returned no parsed results")

    return "\n".join(r["ParsedText"] for r in parsed).strip()
