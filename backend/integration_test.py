"""
CV Extractor -- Full Pipeline Integration Test
Covers: Auth -> Job -> Upload Sign -> Cloudinary Upload ->
        File Register -> Webhook Bypass -> Trigger -> Poll ->
        Rows Verify -> Export

Run from: CvExtractor/cvextractor/
Command:  python integration_test.py
Requires: Django server running on :8000, Celery worker running, Redis running
"""


import os
import sys
import time
import json
import hmac
import hashlib
import requests
from pathlib import Path


# ─── Config ──────────────────────────────────────────────────────────────────

BASE_URL = "http://localhost:8000/api/v1"
TEST_EMAIL = "test@test.com"
TEST_PASSWORD = "Test1234!"


def load_env(path=".env"):
    """Parse .env without python-dotenv dependency."""
    env = {}
    env_path = Path(path)
    if not env_path.exists():
        fail("Setup", f".env file not found at {env_path.absolute()}")
    with open(env_path) as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            k, v = line.split("=", 1)
            env[k.strip()] = v.strip()
    return env


# ─── Helpers ─────────────────────────────────────────────────────────────────

def log(msg, status="INFO"):
    icons = {"INFO": "  →", "OK": "  ✓", "FAIL": "  ✗", "WAIT": "  ⏳", "HEAD": "\n══"}
    print(f"{icons.get(status, '  ?')} {msg}")


def fail(step, detail=""):
    log(f"FAILED: {step}", "FAIL")
    if detail:
        log(detail, "FAIL")
    sys.exit(1)


def ok_or_fail(response, step):
    """Assert response is 2xx and success=true. Return data dict."""
    if not response.ok:
        fail(step, f"HTTP {response.status_code}: {response.text[:400]}")
    body = response.json()
    if not body.get("success"):
        fail(step, f"success=false\n{json.dumps(body, indent=2)[:400]}")
    log(step, "OK")
    return body["data"]


def auth_headers(token):
    return {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}


# ─── Phase 0: Preflight checks ───────────────────────────────────────────────

CLOUDINARY_CLOUD_NAME = ""
CLOUDINARY_API_KEY = ""
CLOUDINARY_API_SECRET = ""
CLOUDINARY_WEBHOOK_SECRET = ""


def load_cloudinary_env():
    """Populate Cloudinary globals from .env. Never hardcode credentials here."""
    global CLOUDINARY_CLOUD_NAME, CLOUDINARY_API_KEY, CLOUDINARY_API_SECRET, CLOUDINARY_WEBHOOK_SECRET
    env = load_env()
    CLOUDINARY_CLOUD_NAME = env.get("CLOUDINARY_CLOUD_NAME", "")
    CLOUDINARY_API_KEY = env.get("CLOUDINARY_API_KEY", "")
    CLOUDINARY_API_SECRET = env.get("CLOUDINARY_API_SECRET", "")
    CLOUDINARY_WEBHOOK_SECRET = env.get("CLOUDINARY_WEBHOOK_SECRET", "")

def phase_preflight():
    log("PHASE 0: PREFLIGHT CHECKS", "HEAD")

    load_cloudinary_env()

    missing = [k for k in ("CLOUDINARY_API_SECRET", "CLOUDINARY_API_KEY",
                            "CLOUDINARY_CLOUD_NAME", "CLOUDINARY_WEBHOOK_SECRET")
               if not globals().get(k)]
    if missing:
        fail("Preflight", f"Missing env vars: {missing}")
    log("Env vars present", "OK")

# ─── Phase 1: Auth ───────────────────────────────────────────────────────────

def phase_auth():
    log("PHASE 1: AUTH", "HEAD")
    r = requests.post(f"{BASE_URL}/auth/login/",
                      json={"email": TEST_EMAIL, "password": TEST_PASSWORD})
    data = ok_or_fail(r, "Login")
    token = data["access_token"]
    log(f"Token acquired (first 20 chars): {token[:20]}...")
    return token


# ─── Phase 2: Create job ─────────────────────────────────────────────────────

def phase_create_job(token):
    log("PHASE 2: CREATE JOB", "HEAD")
    payload = {
        "name": "Integration Test Job",
        "fields": [
            {"key": "name",             "label": "Full Name",        "type": "string", "is_custom": False},
            {"key": "email",            "label": "Email",            "type": "string", "is_custom": False},
            {"key": "skills",           "label": "Skills",           "type": "list",   "is_custom": False},
            {"key": "experience_years", "label": "Experience Years", "type": "number", "is_custom": False},
            {"key": "current_role",     "label": "Current Role",     "type": "string", "is_custom": False},
        ]
    }
    r = requests.post(f"{BASE_URL}/jobs/", json=payload, headers=auth_headers(token))
    data = ok_or_fail(r, "Create job")
    job_id = data["job"]["id"]
    log(f"Job ID: {job_id}")
    log(f"Status: {data['job']['status']}")
    return job_id


# ─── Phase 3: Create test files ──────────────────────────────────────────────

def phase_create_test_files():
    log("PHASE 3: CREATE TEST CV FILES", "HEAD")

    # DOCX -- uses python-docx (already in your stack)
    from docx import Document
    doc = Document()
    doc.add_heading("John Doe", level=0)
    doc.add_paragraph("Email: john.doe@example.com")
    doc.add_paragraph("Phone: +91 98765 43210")
    doc.add_paragraph("Current Role: Senior Backend Engineer")
    doc.add_paragraph("Experience: 5 years")
    doc.add_paragraph("Skills: Python, Django, DRF, PostgreSQL, Redis, Celery, Docker")
    doc.add_paragraph("Location: Mumbai, India")
    docx_path = Path("_test_cv_john.docx")
    doc.save(str(docx_path))
    log(f"Created DOCX: {docx_path}", "OK")

    # JPEG image CV -- uses Pillow (installed in Step 1)
    from PIL import Image, ImageDraw
    img = Image.new("RGB", (900, 700), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    cv_text = (
        "Jane Smith\n"
        "Email: jane.smith@example.com\n"
        "Phone: +91 91234 56789\n"
        "Current Role: Frontend Developer\n"
        "Experience: 3 years\n"
        "Skills: React, Next.js, TypeScript, Node.js, Tailwind CSS\n"
        "Location: Bangalore, India"
    )
    draw.text((60, 60), cv_text, fill=(0, 0, 0))
    img_path = Path("_test_cv_jane.jpg")
    img.save(img_path, "JPEG", quality=90)
    log(f"Created JPEG: {img_path}", "OK")

    return [
        {
            "path": docx_path,
            "file_type": "DOCX",
            "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        },
        {
            "path": img_path,
            "file_type": "IMAGE",
            "mime": "image/jpeg",
        },
    ]


# ─── Phase 4: Sign + upload to Cloudinary + register ─────────────────────────

def phase_upload(token, job_id, test_files):
    log("PHASE 4: UPLOAD (SIGN -> CLOUDINARY -> REGISTER)", "HEAD")
    headers = auth_headers(token)
    registered = []

    for f in test_files:
        path: Path = f["path"]
        file_size = path.stat().st_size
        log(f"Processing: {path.name} ({file_size} bytes)")

        # 4a. Get signed upload params from Django
        sign_payload = {
            "filename": path.name,
            "file_type": f["mime"],
            "file_size": file_size,
        }
        r = requests.post(
            f"{BASE_URL}/jobs/{job_id}/upload/sign/",
            json=sign_payload,
            headers=headers,
        )
        sign_data = ok_or_fail(r, f"Sign upload: {path.name}")

        upload_url = sign_data["upload_url"]
        upload_params = sign_data["upload_params"]
        cloudinary_public_id = sign_data["cloudinary_public_id"]

        # 4b. Upload directly to Cloudinary
        log(f"Uploading {path.name} to Cloudinary...")
        with open(path, "rb") as fh:
            form_data = {k: str(v) for k, v in upload_params.items()}
            cld_response = requests.post(
                upload_url,
                data=form_data,
                files={"file": (path.name, fh, f["mime"])},
            )

        if not cld_response.ok:
            fail(f"Cloudinary upload: {path.name}",
                 f"HTTP {cld_response.status_code}: {cld_response.text[:300]}")
        cld_data = cld_response.json()
        log(f"Cloudinary accepted: {path.name} -> {cld_data.get('secure_url', 'N/A')}", "OK")

        registered.append({
            "cloudinary_public_id": cloudinary_public_id,
            "original_filename": path.name,
            "file_type": f["file_type"],
            # Needed for webhook bypass payload
            "secure_url": cld_data.get("secure_url"),
            "bytes": cld_data.get("bytes"),
            "format": cld_data.get("format"),
        })

    # 4c. Register all files with Django in one batch call
    reg_payload = {
        "files": [
            {
                "cloudinary_public_id": f["cloudinary_public_id"],
                "original_filename": f["original_filename"],
                "file_type": f["file_type"],
            }
            for f in registered
        ]
    }
    r = requests.post(f"{BASE_URL}/jobs/{job_id}/files/", json=reg_payload, headers=headers)
    reg_data = ok_or_fail(r, "Register files (batch)")
    log(f"Registered: {reg_data['registered']} files")
    log(f"Job status now: {reg_data['job_status']}")

    if reg_data["job_status"] != "QUEUED":
        fail("File registration", f"Expected QUEUED, got {reg_data['job_status']}")

    return registered


# ─── Phase 5: Webhook bypass ─────────────────────────────────────────────────

def phase_webhook_bypass(registered_files):
    log("PHASE 5: WEBHOOK VERIFY (signed Cloudinary callback)", "HEAD")
    log("Calls /webhooks/cloudinary/ with a signature matching the real handler.")

    for f in registered_files:
        payload = {
            "public_id": f["cloudinary_public_id"],
            "secure_url": f["secure_url"],
            "format": f["format"],
            "bytes": f["bytes"],
            "notification_type": "upload",
        }
        body_str = json.dumps(payload, separators=(",", ":"))
        timestamp = str(int(time.time()))

        to_sign = body_str + timestamp + CLOUDINARY_WEBHOOK_SECRET
        signature = hashlib.sha1(to_sign.encode("utf-8")).hexdigest()

        r = requests.post(
            f"{BASE_URL}/webhooks/cloudinary/",
            data=body_str.encode("utf-8"),
            headers={
                "Content-Type": "application/json",
                "X-Cld-Signature": signature,
                "X-Cld-Timestamp": timestamp,
            },
        )

        if r.status_code != 200:
            fail(
                f"Webhook verify: {f['original_filename']}",
                f"HTTP {r.status_code}: {r.text[:300]}",
            )
        log(f"Webhook accepted: {f['original_filename']}", "OK")

    log("All files marked VERIFIED via signed webhook callback")

# ─── Phase 6: Trigger extraction ─────────────────────────────────────────────

def phase_trigger(token, job_id):
    log("PHASE 6: TRIGGER EXTRACTION", "HEAD")
    r = requests.post(
        f"{BASE_URL}/jobs/{job_id}/process/",
        headers=auth_headers(token),
    )
    data = ok_or_fail(r, "Trigger process")
    log(f"Job status: {data['status']}")
    log(f"Total files queued: {data['total_files']}")

    if data["status"] != "PROCESSING":
        fail("Trigger", f"Expected PROCESSING, got {data['status']}")


# ─── Phase 7: Poll until terminal state ──────────────────────────────────────

def phase_poll(token, job_id, timeout_seconds=300):
    log("PHASE 7: POLLING STATUS", "HEAD")
    log(f"Polling every 5s, timeout: {timeout_seconds}s")
    log("Celery worker must be running in a separate terminal.")

    headers = auth_headers(token)
    start = time.time()

    while time.time() - start < timeout_seconds:
        r = requests.get(f"{BASE_URL}/jobs/{job_id}/status/", headers=headers)
        data = ok_or_fail(r, "Poll status")

        status = data["status"]
        done = data["done_files"]
        failed = data["failed_files"]
        total = data["total_files"]
        elapsed = int(time.time() - start)

        log(f"[{elapsed}s] {status} | Done: {done}/{total} | Failed: {failed}", "WAIT")

        if status == "COMPLETE":
            log(f"Job COMPLETE. All {total} files extracted successfully.", "OK")
            return "COMPLETE"
        if status == "PARTIAL":
            log(f"Job PARTIAL. {done} done, {failed} failed.", "OK")
            return "PARTIAL"
        if status == "FAILED":
            log(f"Job FAILED. All {total} files failed extraction.", "FAIL")
            return "FAILED"

        time.sleep(5)

    fail("Polling", f"Timed out after {timeout_seconds}s. Check Celery worker logs for errors.")


# ─── Phase 8: Verify rows ────────────────────────────────────────────────────

def phase_verify_rows(token, job_id):
    log("PHASE 8: VERIFY EXTRACTED ROWS", "HEAD")
    r = requests.get(f"{BASE_URL}/jobs/{job_id}/rows/", headers=auth_headers(token))
    data = ok_or_fail(r, "Fetch rows")

    rows = data["rows"]
    log(f"Total rows returned: {len(rows)}")

    all_fields_null = []

    for row in rows:
        log(f"\n  File: {row['original_filename']} | Status: {row['extraction_status']}")
        extracted = row["data"] or {}
        non_null = {k: v for k, v in extracted.items() if v is not None}
        null_fields = [k for k, v in extracted.items() if v is None]

        for k, v in extracted.items():
            marker = "NULL" if v is None else "OK  "
            log(f"    [{marker}] {k}: {v}")

        if not non_null:
            all_fields_null.append(row["original_filename"])

    if all_fields_null:
        log(f"\nWARNING: These files had ALL fields null:", "FAIL")
        for name in all_fields_null:
            log(f"  - {name}", "FAIL")
        log("Check Celery logs for extraction errors on these files.")
    else:
        log("All rows have at least some extracted data.", "OK")

    return rows


# ─── Phase 9: Export ─────────────────────────────────────────────────────────

def phase_export(token, job_id):
    log("PHASE 9: EXPORT", "HEAD")
    headers = auth_headers(token)
    # Remove Content-Type for file download (no body being sent)
    headers.pop("Content-Type", None)

    r = requests.get(f"{BASE_URL}/jobs/{job_id}/export/?export_format=xlsx", headers=headers)

    if not r.ok:
        fail("Export", f"HTTP {r.status_code}: {r.text[:300]}")

    content_type = r.headers.get("Content-Type", "")
    if "spreadsheet" not in content_type and "octet-stream" not in content_type:
        fail("Export", f"Unexpected Content-Type: {content_type}\nBody: {r.text[:200]}")

    export_path = Path("_test_export_output.xlsx")
    export_path.write_bytes(r.content)
    log(f"Export saved: {export_path} ({len(r.content)} bytes)", "OK")
    log("Open _test_export_output.xlsx and verify columns and data match what was extracted.")


# ─── Cleanup ─────────────────────────────────────────────────────────────────

def cleanup():
    for name in ["_test_cv_john.docx", "_test_cv_jane.jpg"]:
        p = Path(name)
        if p.exists():
            p.unlink()
    log("Test files cleaned up", "OK")


# ─── Main ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("\n" + "=" * 60)
    print("  CV EXTRACTOR -- FULL PIPELINE INTEGRATION TEST")
    print("=" * 60)

    final_status = None
    try:
        phase_preflight()
        token   = phase_auth()
        job_id  = phase_create_job(token)
        files   = phase_create_test_files()
        reg     = phase_upload(token, job_id, files)
        phase_webhook_bypass(reg)
        phase_trigger(token, job_id)
        final_status = phase_poll(token, job_id)

        if final_status in ("COMPLETE", "PARTIAL"):
            phase_verify_rows(token, job_id)
            phase_export(token, job_id)
        else:
            log("Skipping rows + export (job FAILED)", "FAIL")

    except KeyboardInterrupt:
        log("Interrupted by user", "FAIL")
    finally:
        cleanup()

    print("\n" + "=" * 60)
    if final_status in ("COMPLETE", "PARTIAL"):
        print(f"  RESULT: {final_status} -- Pipeline is working.")
    elif final_status == "FAILED":
        print("  RESULT: FAILED -- Check Celery worker logs.")
    print("=" * 60 + "\n")